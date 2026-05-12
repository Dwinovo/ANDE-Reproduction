"""Generate the two classroom presentation Word documents.

Builds, at the repository root:
  * `ANDE复现_课堂汇报.docx`        — 复现部分
  * `ANDE扩展实验_课堂汇报.docx`    — 扩展实验部分

Each document is self-contained (the extension doc opens with a short
recap of what ANDE is and what we reproduced, so it can be presented
without the reproduction doc being shown first).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

REPO = Path(__file__).resolve().parent.parent
FIG = REPO / "docs" / "figures"
OUTPUT_REPRO = REPO / "ANDE复现_课堂汇报.docx"
OUTPUT_EXT = REPO / "ANDE扩展实验_课堂汇报.docx"

EAST_ASIA_FONT = "Microsoft YaHei"
LATIN_FONT = "Calibri"
MONO_FONT = "Consolas"


# =====================================================================
# Style helpers
# =====================================================================

def _set_run_font(run, *, size_pt: float = 11, bold: bool = False,
                  italic: bool = False, color: RGBColor | None = None,
                  mono: bool = False) -> None:
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    run.font.name = MONO_FONT if mono else LATIN_FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), MONO_FONT if mono else EAST_ASIA_FONT)
    rFonts.set(qn("w:ascii"), MONO_FONT if mono else LATIN_FONT)
    rFonts.set(qn("w:hAnsi"), MONO_FONT if mono else LATIN_FONT)


def _make_normal_style(doc: Document) -> None:
    from docx.oxml import OxmlElement
    style = doc.styles["Normal"]
    style.font.name = LATIN_FONT
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    rFonts.set(qn("w:ascii"), LATIN_FONT)
    rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    rPr.append(rFonts)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    style_map = {0: ("Title", 26, True), 1: ("Heading 1", 18, True),
                 2: ("Heading 2", 15, True), 3: ("Heading 3", 13, True)}
    style_name, size, bold = style_map.get(level, ("Heading 4", 12, True))
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _set_run_font(run, size_pt=size, bold=bold,
                  color=RGBColor(0x1F, 0x3A, 0x68) if level <= 1 else None)


def add_para(doc: Document, text: str, *, size: float = 11,
             bold: bool = False, italic: bool = False, indent_cm: float = 0,
             align: int | None = None) -> None:
    p = doc.add_paragraph()
    if indent_cm:
        p.paragraph_format.first_line_indent = Cm(indent_cm)
    p.paragraph_format.space_after = Pt(4)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    _set_run_font(run, size_pt=size, bold=bold, italic=italic)


def add_bullet(doc: Document, text: str, *, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    _set_run_font(run, size_pt=11)


def add_code_block(doc: Document, text: str, *, size: float = 9.5) -> None:
    from docx.oxml import OxmlElement

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    cell.width = Cm(15.5)
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    tc_pr.append(shd)
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "BFBFBF")
        tcBorders.append(b)
    tc_pr.append(tcBorders)

    cell.text = ""
    for i, line in enumerate(text.rstrip().split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line)
        _set_run_font(run, size_pt=size, mono=True)


def add_callout(doc: Document, title: str, body: str,
                color_hex: str = "FFF4DC") -> None:
    from docx.oxml import OxmlElement

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    cell.width = Cm(15.5)
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:color"), "E1B45A")
        tcBorders.append(b)
    tc_pr.append(tcBorders)

    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    _set_run_font(run, size_pt=11, bold=True, color=RGBColor(0x99, 0x57, 0x00))
    for line in body.strip().split("\n"):
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(line)
        _set_run_font(r2, size_pt=10.5)


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              *, col_widths_cm: list[float] | None = None,
              align: list[int] | None = None,
              size: float = 10) -> None:
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)

    def _fill_cell(cell, text: str, *, bold: bool = False, header: bool = False,
                   alignment: int | None = None) -> None:
        from docx.oxml import OxmlElement

        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""
        if header:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "1F3A68")
            tc_pr.append(shd)
        p = cell.paragraphs[0]
        if alignment is not None:
            p.alignment = alignment
        remaining = text
        while remaining:
            if "**" in remaining:
                pre, sep, rest = remaining.partition("**")
                if pre:
                    r = p.add_run(pre)
                    _set_run_font(r, size_pt=size, bold=bold or header,
                                  color=RGBColor(0xFF, 0xFF, 0xFF) if header else None)
                bold_text, _, after = rest.partition("**")
                r = p.add_run(bold_text)
                _set_run_font(r, size_pt=size, bold=True,
                              color=RGBColor(0xFF, 0xFF, 0xFF) if header else None)
                remaining = after
            else:
                r = p.add_run(remaining)
                _set_run_font(r, size_pt=size, bold=bold or header,
                              color=RGBColor(0xFF, 0xFF, 0xFF) if header else None)
                remaining = ""

    for i, h in enumerate(headers):
        _fill_cell(tbl.rows[0].cells[i], h, header=True,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for r_i, row in enumerate(rows):
        for c_i, cell_text in enumerate(row):
            alignment = None
            if align and c_i < len(align):
                alignment = align[c_i]
            _fill_cell(tbl.rows[r_i + 1].cells[c_i], cell_text,
                       alignment=alignment)


def add_image(doc: Document, path: Path, *, width_cm: float = 14,
              caption: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(8)
        crun = cap.add_run(caption)
        _set_run_font(crun, size_pt=9.5, italic=True,
                      color=RGBColor(0x60, 0x60, 0x60))


# =====================================================================
# Covers
# =====================================================================

def cover_repro(doc: Document) -> None:
    add_heading(doc, "ANDE 论文复现", level=0)
    add_para(doc, "课堂汇报（第一部分）· 三轮弯路的真实故事",
             size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "")
    add_para(doc, "论文：Deng et al., 《ANDE: Detect the Anonymity Web Traffic With "
                  "Comprehensive Model》, IEEE TNSM 2024",
             size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    add_para(doc, "实验环境：AutoDL · NVIDIA RTX 5090 32GB · "
                  "PyTorch 2.8.0 + cu128",
             size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    add_para(doc, "时间：2026-05-07 ~ 2026-05-09",
             size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    add_para(doc, "")
    add_callout(
        doc,
        "一句话总结",
        "本汇报 = 把一篇 IEEE 期刊论文从论文还原成可运行代码 + 数据 + 完整复现结果。\n"
        "整个过程走了 3 轮弯路，最终 ANDE 8100B / 14 类 accuracy = 0.9458（论文 0.9820，差 3.6 pp）。\n"
        "前两轮失败的过程，比最终结果本身更有讨论价值。",
    )


def cover_extension(doc: Document) -> None:
    add_heading(doc, "ANDE 扩展实验", level=0)
    add_para(doc, "课堂汇报（第二部分）· 在干净复现的基础上继续追问",
             size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "")
    add_para(doc, "论文：Deng et al., 《ANDE: Detect the Anonymity Web Traffic With "
                  "Comprehensive Model》, IEEE TNSM 2024",
             size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    add_para(doc, "实验环境：AutoDL · NVIDIA RTX 5090 32GB",
             size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    add_para(doc, "时间：2026-05-10 ~ 2026-05-12",
             size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    add_para(doc, "")
    add_callout(
        doc,
        "本场汇报的三个核心问题",
        "Q1：论文把字节硬 reshape 成 90×90 二维方阵——这真的合理吗？\n"
        "Q2：784 / 4096 / 8100 三种固定长度太朴素，能不能让模型自己选关键字节段？\n"
        "Q3：攻击者用 padding / 延迟 / 流量整形，能不能把 ANDE 骗瞎？\n"
        "全部用真实实验回答，并对每个问题做了 single-seed 探索 + multi-seed 上规模验证。",
    )


# =====================================================================
# Reproduction chapters
# =====================================================================

def chap_repro_intro(doc: Document) -> None:
    add_heading(doc, "一、我们要研究什么", level=1)

    add_heading(doc, "1.1 现实背景：匿名网络流量检测", level=2)
    add_para(
        doc,
        "Tor、I2P 这类匿名网络让用户隐藏自己的真实 IP。这对个人隐私是好事，"
        "但对网络运营商来说就是个大麻烦——攻击者也躲在里面。问题是，加密以后"
        "网络包的 payload 已经不可读，传统基于关键字匹配的 DPI（深度包检测）"
        "失效了。",
        indent_cm=0.74,
    )
    add_para(
        doc,
        "ANDE 这篇论文要回答的就是：**只看密文流量的「形状」，能不能识别"
        "用户行为？** 比如这个 session 是在用 Tor 看视频，还是在用 Tor 聊天？",
        indent_cm=0.74,
    )

    add_heading(doc, "1.2 本项目的两个任务", level=2)
    add_table(
        doc,
        headers=["任务名", "类别数", "说明"],
        rows=[
            ["Binary2", "2", "Tor / NonTor 二分类（简单基准）"],
            ["Behavior14", "**14**", "**7 种行为 × {Tor, NonTor}**——这才是真正难的"],
        ],
        col_widths_cm=[3.5, 2, 10],
        align=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, None],
    )
    add_para(doc, "14 个具体类别的标签编号：", size=10.5)
    add_code_block(
        doc,
        "browsing-nontor = 0      browsing-tor    = 1\n"
        "chat-nontor     = 2      chat-tor        = 3\n"
        "email-nontor    = 4      email-tor       = 5\n"
        "ft-nontor       = 6      ft-tor          = 7\n"
        "p2p-nontor      = 8      p2p-tor         = 9\n"
        "streaming-nontor= 10     streaming-tor   = 11\n"
        "voip-nontor     = 12     voip-tor        = 13",
    )


def chap_sample_definition(doc: Document) -> None:
    add_heading(doc, "二、先讲清楚：一个「样本」到底长什么样", level=1)
    add_para(
        doc,
        "整个复现里最容易踩坑的地方，就是「pcap、packet、session」三个单位混用。"
        "这里先用一个真实的 pcap 文件举例说明。",
    )

    add_heading(doc, "2.1 pcap：原始抓包文件", level=2)
    add_para(doc, "以本项目实际用到的一个 pcap 为例：")
    add_code_block(doc, "文件名：p2p_multipleSpeed2-1.pcap\n"
                        "标签：p2p-nontor （非 Tor 的 P2P 下载）\n"
                        "文件大小：~2.4 GB\n"
                        "包含 packet 数：~1600 万")

    add_heading(doc, "2.2 packet：一个网络包", level=2)
    add_code_block(
        doc,
        "一个 packet 包含：\n"
        "  时间戳 (timestamp)\n"
        "  源 IP / 源端口\n"
        "  目的 IP / 目的端口\n"
        "  协议号 (TCP=6 / UDP=17 / ICMP=1 ...)\n"
        "  TCP flags (SYN/ACK/FIN/PSH/RST/URG)\n"
        "  payload（加密后的字节）\n"
        "  包长度",
    )

    add_heading(doc, "2.3 session：一条连接（这才是我们的样本单位）", level=2)
    add_para(
        doc,
        "把同属一条 TCP/UDP 连接的所有 packet 聚在一起，就是一个 session。"
        "聚合方法是「五元组」：",
    )
    add_code_block(doc, "session_key = (源IP, 源端口, 目的IP, 目的端口, 协议)")
    add_para(
        doc,
        "为了不区分发起方和响应方，把 (源IP, 源端口) 和 (目的IP, 目的端口) "
        "按字典序排个序，这样回包和发包归到同一个 session。",
    )
    add_para(doc, "上面那个 pcap 切完之后，得到很多 session，比如：")
    add_code_block(
        doc,
        "session_id = p2p_multipleSpeed2-1__10.152.152.11_55357_103.52.253.34_27965_6\n"
        "              └─── pcap 名 ───┘   └── src ──┘  └────── dst ──────┘  └ proto\n"
        "\n"
        "label  = p2p-nontor\n"
        "包数   = 247 个 packet\n"
        "时长   = 123.4 秒",
    )
    add_para(
        doc,
        "这个 pcap 一共切出 **3265 条有效 session**（保留 ≥ 3 个包的）。"
        "也就是说，1 个 pcap → 3265 个训练样本。",
    )

    if (FIG / "classroom_sample_definition.png").exists():
        add_image(doc, FIG / "classroom_sample_definition.png", width_cm=14,
                  caption="图 2-1 · pcap / packet / session 三层单位关系")

    add_heading(doc, "2.4 整个项目的数据规模", level=2)
    add_table(
        doc,
        headers=["数据集", "pcap 数", "体积"],
        rows=[
            ["ISCXTor2016 / Tor", "50", "12.5 GB"],
            ["ISCXTor2016 / NonTor", "44", "10.6 GB"],
            ["darknet-2020 / Tor", "60", "9.9 GB"],
            ["**合计**", "**154**", "**33.0 GB**"],
        ],
        col_widths_cm=[6, 3, 4],
        align=[None, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
    )
    add_para(doc, "预处理产物：")
    add_bullet(doc, "Algorithm 1 输出（含 ≥ 3 包的 session）：54,460 条")
    add_bullet(doc, "Algorithm 2 输出（含 ≥ 10 包的 session）：24,995 条")
    add_bullet(doc, "最终训练用：24,995（两者 inner join），train/test = 19,996 / 4,999")


def chap_paper_method(doc: Document) -> None:
    add_heading(doc, "三、ANDE 论文的方法：两个 Algorithm + 一个双分支模型", level=1)

    if (FIG / "classroom_algorithm_ande_overview.png").exists():
        add_image(doc, FIG / "classroom_algorithm_ande_overview.png", width_cm=15.5,
                  caption="图 3-1 · Algorithm 1 + Algorithm 2 + ANDE 模型的整体管道")

    # Algorithm 1
    add_heading(doc, "3.1 Algorithm 1：把 session 变成「灰度图」", level=2)
    add_para(doc, "目标：让 CNN 能直接看 session 的原始字节。具体到代码就 5 步：")
    add_code_block(
        doc,
        "for each session in pcap:\n"
        "    1) 读取这条 session 内的所有 packet\n"
        "    2) 取每个 packet 的原始 bytes\n"
        "    3) 按时间顺序拼接成一条长字节流\n"
        "    4) 截断或零补齐到固定长度 γ\n"
        "    5) reshape 成 √γ × √γ 的灰度图",
    )
    add_para(doc, "论文用了三个长度，对应三种图：")
    add_table(
        doc,
        headers=["输入长度 γ (bytes)", "灰度图 shape", "代表性"],
        rows=[
            ["784", "28 × 28", "和 MNIST 一样大"],
            ["4096", "64 × 64", "覆盖 ~82% 的 session"],
            ["8100", "90 × 90", "覆盖更长 session"],
        ],
        col_widths_cm=[4.5, 4, 7],
        align=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, None],
    )

    add_para(doc, "另外要做一步重要的「匿名化」：把每个 packet 的以下字段抹零，"
                  "防止模型「作弊」靠 IP/MAC 记忆：")
    add_code_block(
        doc,
        "Ethernet src / dst MAC   → 00:00:00:00:00:00\n"
        "IP       src / dst       → 0.0.0.0\n"
        "TCP/UDP  sport / dport   → 0",
    )
    add_para(doc, "样本看上去是这样（同一行是同一行为的 14 个类别）：")
    if (FIG / "sample_images.png").exists():
        add_image(doc, FIG / "sample_images.png", width_cm=15.5,
                  caption="图 3-2 · 14 个类别各取一个 session 的灰度图")

    # Algorithm 2
    add_heading(doc, "3.2 Algorithm 2：26 维统计特征", level=2)
    add_para(
        doc,
        "光看字节图还不够。论文又针对每条 session 算了 26 个手工统计特征，"
        "这些特征更接近传统流量分析的视角（包大小、时延、TCP flag 比例等）。"
        "26 维特征清单：",
    )
    add_table(
        doc,
        headers=["类别", "数量", "具体特征"],
        rows=[
            ["TCP flag 占比", "6",
             "Avg_{syn, urg, fin, ack, psh, rst}_flag"],
            ["协议占比", "4",
             "Avg_{DNS, TCP, UDP, ICMP}_pkt"],
            ["时间", "5",
             "Duration / Avg/Min/Max/StDev_deltas_time"],
            ["包长度", "4",
             "Avg/Min/Max/StDev_Pkts_length"],
            ["小包比例", "1", "Avg_small_loadings_pkt (<32B)"],
            ["payload", "4", "Avg/Min/Max/StDev_payload"],
            ["DNS / TCP 比", "1", "Avg_DNS_over_TCP"],
            ["包总数", "1", "num_packets"],
        ],
        col_widths_cm=[3.5, 1.5, 10.5],
        align=[None, WD_ALIGN_PARAGRAPH.CENTER, None],
    )
    add_para(
        doc,
        "在所有 session 上算完后做 z-score 标准化：每个特征减去全集均值、再除以全集标准差。",
        size=10.5,
    )
    add_para(doc, "举例：一个真实 session 的部分特征（标准化前 / 后）：")
    add_table(
        doc,
        headers=["特征名", "原值", "z-score 标准化后"],
        rows=[
            ["num_packets", "247", "-0.0211"],
            ["Duration_window_flow", "123.4 s", "-0.0149"],
            ["Avg_Pkts_length", "512.3 B", "-0.3372"],
            ["Avg_syn_flag", "4 / 247 = 0.016", "(标准化值)"],
            ["Avg_ack_flag", "0.95", "(标准化值)"],
        ],
        col_widths_cm=[5.5, 4, 5],
        align=[None, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
    )

    # ANDE model
    add_heading(doc, "3.3 ANDE 模型：双分支 + 融合", level=2)
    add_para(doc, "ANDE 把 Algorithm 1 的灰度图 + Algorithm 2 的 26 维"
                  "向量分别送进两条支路，最后融合分类：")
    add_code_block(
        doc,
        "                   ┌──────────────────────┐\n"
        " 灰度图  (1×90×90) │  SE-ResNet-18         │  256 维\n"
        "  ───────────────► │  (image backbone)     │ ────────┐\n"
        "                   └──────────────────────┘         │\n"
        "                                                    ▼\n"
        "                                              ┌─────────────┐\n"
        "                                              │ concat       │\n"
        "                                              │  → 265 维   │\n"
        "                   ┌──────────────────────┐ ▲│  → 100 → 30 │ → 14 类\n"
        " 26 维 stats      │  StatMLP 26→18→9     │ ││             │\n"
        "  ───────────────► │  (stat  backbone)    │ │└─────────────┘\n"
        "                   └──────────────────────┘ │\n"
        "                              9 维 ─────────┘",
    )
    add_para(doc, "关键设计点：")
    add_bullet(doc, "SE-ResNet-18：标准 ResNet-18 但通道减半，每个 BasicBlock "
                    "里都插了一个 SEBlock 通道注意力")
    add_bullet(doc, "SEBlock：全局平均池化 → FC 降维 → ReLU → FC 升维 → Sigmoid，"
                    "得到每个通道的权重再乘回特征图")
    add_bullet(doc, "StatMLP：26 → 18 → 9，把统计特征压成 9 维")
    add_bullet(doc, "FusionHead：265 → 100 → 30 → C 的三层 MLP 出 logits")
    add_para(doc, "整个网络约 **285 万参数**，相比 ResNet-50 这种动辄几千万的"
                  "网络已经算「轻量」。")


def chap_three_rounds(doc: Document) -> None:
    add_heading(doc, "四、复现走过的三轮弯路（重点章节）", level=1)
    add_para(
        doc,
        "我们并不是一上来就拿到漂亮的 0.9458 结果。这是经过 3 轮试错才"
        "「干净地」复现出来的。三轮的唯一区别其实就是一句话：「统计特征到底"
        "按什么单位算？数据集按什么单位划分？」",
    )
    add_table(
        doc,
        headers=["轮次", "image 单位", "stats 单位", "split 单位",
                 "ANDE 8100 acc", "结论"],
        rows=[
            ["Round 1", "session", "**pcap**", "session", "**0.9908**",
             "❌ 数据泄漏，看着漂亮但是假象"],
            ["Round 2", "session", "pcap", "**pcap**", "0.6579",
             "⚠️ 堵住泄漏但过度修正了"],
            ["**Round 3 (final)**", "session", "**session**", "session",
             "**0.9458**", "✅ 干净复现"],
        ],
        col_widths_cm=[2.5, 1.8, 1.8, 1.8, 2.2, 5.5],
        align=[WD_ALIGN_PARAGRAPH.CENTER] * 5 + [None],
    )

    if (FIG / "experiment_process_rounds.png").exists():
        add_image(doc, FIG / "experiment_process_rounds.png", width_cm=15,
                  caption="图 4-1 · 三轮实验在所有方法上的 accuracy 对照")

    # ----- Round 1 -----
    add_heading(doc, "4.1 Round 1：按论文字面意思实现，结果「太好了」", level=2)
    add_para(doc, "**怎么做的：**")
    add_code_block(
        doc,
        "Algorithm 1：按 session 切图  → image 是 session 级\n"
        "Algorithm 2：按整个 pcap 算   → stats  是 pcap 级（论文伪代码字面理解）\n"
        "Split：按 session 随机 8:2  → split  是 session 级",
    )
    add_para(doc, "**问题出在哪？** 还是看那个 p2p_multipleSpeed2-1.pcap 的例子。"
                  "它有 3265 条 session，但 stats 只算了一遍：")
    add_code_block(
        doc,
        "整个 p2p_multipleSpeed2-1.pcap 算出来一份 26 维向量 stats_p2p\n"
        "      ↓\n"
        "S1     image_S1   + stats_p2p + p2p-nontor\n"
        "S2     image_S2   + stats_p2p + p2p-nontor    ←── stats 完全一样！\n"
        "S3     image_S3   + stats_p2p + p2p-nontor\n"
        " ...\n"
        "S3265  image_S3265+ stats_p2p + p2p-nontor",
    )

    add_para(doc, "session-level 8:2 split 之后：")
    add_table(
        doc,
        headers=["split", "session_id 示例", "stats", "label"],
        rows=[
            ["train", "p2p_multipleSpeed2-1__...55357_...27965_6",
             "stats_p2p", "p2p-nontor"],
            ["train", "p2p_multipleSpeed2-1__...37070_...8999_6",
             "stats_p2p ← 同一个", "p2p-nontor"],
            ["test", "p2p_multipleSpeed2-1__...37069_...38553_6",
             "stats_p2p ← 又是它", "p2p-nontor"],
            ["test", "p2p_multipleSpeed2-1__...9100_...57558_17",
             "stats_p2p ← 还是它", "p2p-nontor"],
        ],
        col_widths_cm=[1.5, 7.5, 3.5, 3],
        align=[WD_ALIGN_PARAGRAPH.CENTER, None,
               WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
    )
    add_para(doc, "**这就是数据泄漏**：训练集的 2595 条 session 和测试集的 670 条 "
                  "session 共用同一份 stats。决策树只要按 stats 阈值就能秒分对。")

    if (FIG / "classroom_round1_leakage.png").exists():
        add_image(doc, FIG / "classroom_round1_leakage.png", width_cm=15,
                  caption="图 4-2 · Round 1 数据泄漏示意")

    add_para(doc, "**Round 1 的结果（8100B / 14 类）：**")
    add_table(
        doc,
        headers=["Method", "Accuracy", "异常"],
        rows=[
            ["DT", "**0.9999**", "🚨 决策树达到 99.99%"],
            ["RF", "**1.0000**", "🚨 100%"],
            ["XGB", "**1.0000**", "🚨 100%"],
            ["CNN1D", "0.9467", "正常"],
            ["ResNet-18", "0.9567", "正常"],
            ["ANDE-no-SE", "0.9916", "略高"],
            ["**ANDE**", "**0.9908 ± 0.0018**", "看上去美得不真实"],
        ],
        col_widths_cm=[3.5, 4, 8],
        align=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, None],
    )
    add_callout(
        doc,
        "Round 1 的报警信号是什么？",
        "不是 ANDE 太高，而是 DT / RF / XGB 全部 100%。这种简单模型不可能在 14 类"
        "上达到满分，除非「答案已经写在输入里」。这个信号让我们意识到「stats 这一路"
        "肯定哪里漏了」。",
    )

    # ----- Round 2 -----
    add_heading(doc, "4.2 Round 2：堵泄漏，但是堵过头了", level=2)
    add_para(doc, "**怎么做的：** 把 split 改成 pcap-level，让 train 的 pcap "
                  "和 test 的 pcap 完全不重合。")
    add_code_block(
        doc,
        "Algorithm 1：按 session 切图\n"
        "Algorithm 2：按整个 pcap 算\n"
        "Split：按 **pcap 文件** 8:2 划分 ← 同一 pcap 的所有 session 一起进 train 或 test",
    )
    add_para(doc, "**举个具体例子（ft-nontor 类别）：**")
    add_table(
        doc,
        headers=["pcap 文件", "session 数", "Round 2 split"],
        rows=[
            ["SFTP_filetransfer.pcap", "12", "**全部进 train**"],
            ["FTP_filetransfer.pcap", "1544", "**全部进 test**"],
        ],
        col_widths_cm=[6, 3, 6],
        align=[None, WD_ALIGN_PARAGRAPH.CENTER, None],
    )
    add_para(
        doc,
        "也就是说，模型在「**12 条 SFTP**」上学到了什么叫 ft-nontor，"
        "然后被要求去识别「**1544 条 FTP**」。这两个协议虽然都叫文件传输，"
        "但流量特征差异巨大。任务变成了「跨 pcap 文件泛化」，比论文原意更严格。",
    )

    if (FIG / "classroom_round2_concrete.png").exists():
        add_image(doc, FIG / "classroom_round2_concrete.png", width_cm=15,
                  caption="图 4-3 · Round 2 的过度修正")

    add_para(doc, "**Round 2 结果（8100B / 14 类）：**")
    add_table(
        doc,
        headers=["Method", "Round 1", "Round 2", "评价"],
        rows=[
            ["DT", "0.9999", "0.7404", "断崖下降"],
            ["RF", "1.0000", "0.7481", "断崖下降"],
            ["XGB", "1.0000", "0.6157", "断崖下降"],
            ["CNN1D", "0.9467", "0.5559", ""],
            ["ResNet-18", "0.9567", "0.5767", ""],
            ["ANDE-no-SE", "0.9916", "0.6744", ""],
            ["**ANDE**", "**0.9908**", "**0.6579**", "❌ 不能作为复现"],
        ],
        col_widths_cm=[3.5, 3, 3, 6],
        align=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
               WD_ALIGN_PARAGRAPH.CENTER, None],
    )
    add_callout(
        doc,
        "Round 2 的教训",
        "数据泄漏修复要「对症下药」。这次我们用 pcap 级 split 解决了 stats 共享问题，"
        "但代价是把任务定义改成了「跨 pcap 文件泛化」，比论文设定更难。"
        "0.6579 不能拿去和论文 0.9820 比。",
        color_hex="FFE8E8",
    )

    # ----- Round 3 -----
    add_heading(doc, "4.3 Round 3：重读论文，找到真正答案", level=2)
    add_para(doc, "Round 2 之后我们停下来重新读论文。三个证据指向同一个结论：")
    add_para(doc, "**证据 1：** 论文 Table II 每一行特征定义里都写着 "
                  "「in one session」/「within a packet window」。比如：")
    add_code_block(
        doc,
        "num_packets:  \"Number of packets in **one session**\"\n"
        "Avg_payload:  \"average size of payload **in one session**\"\n"
        "Duration_window_flow:  \"duration of the **packet window**\"",
    )
    add_para(doc, "**证据 2：** 论文报告样本数是「50,905」。我们 154 个 pcap、"
                  "54,460 个 session，只有按 session 算才能匹配这个量级；"
                  "按 pcap 算只有 154 行。")
    add_para(doc, "**证据 3：** Algorithm 1 第 5 行写着 「save session to folders」——"
                  "也就是说，**Algorithm 2 里的 `*.pcap` 其实指 Algorithm 1 输出的「per-session 小 pcap」**，"
                  "不是原始大 pcap。我们被论文伪代码误导了。")

    add_callout(
        doc,
        "Round 3 的关键洞察",
        "「Algorithm 2 中 for *.pcap in folder do 这句话，里面的 *.pcap 指的是"
        "Algorithm 1 输出的 per-session pcap 文件，不是原始抓包。」\n"
        "→ 所以 26 维统计特征应该 **每条 session 算一份**，每条 session 都有自己独有的 stats。",
        color_hex="E0F4E0",
    )

    add_para(doc, "**Round 3 怎么做的：**")
    add_code_block(
        doc,
        "Algorithm 1：按 session 切图\n"
        "Algorithm 2：按 session 算 stats  ← 每条 session 算一份独立的 26 维向量\n"
        "Split：按 session 随机 8:2",
    )
    add_para(doc, "还是看那个 p2p_multipleSpeed2-1.pcap 的同样几条 session，"
                  "现在每条都有自己独立的 stats：")
    add_table(
        doc,
        headers=["split", "session_id", "Duration", "Avg_Pkts_length",
                 "Avg_payload"],
        rows=[
            ["train", "...55357_...27965_6", "-0.0211", "-0.3372", "-0.3372"],
            ["train", "...37070_...8999_6", "-0.0149", "-0.5858", "-0.5874"],
            ["test", "...37069_...38553_6", "-0.0132", "+1.1500", "+1.1540"],
            ["test", "...9100_...57558_17", "+0.1053", "-0.8545", "-0.8056"],
        ],
        col_widths_cm=[1.5, 5.5, 2.5, 3, 3],
        align=[WD_ALIGN_PARAGRAPH.CENTER, None,
               WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
               WD_ALIGN_PARAGRAPH.CENTER],
    )
    add_para(doc, "可以看到：同一 pcap 内的 session，stats 不再完全一样。"
                  "session-level split 也就不再泄漏。", size=10.5)

    if (FIG / "classroom_round3_correct.png").exists():
        add_image(doc, FIG / "classroom_round3_correct.png", width_cm=15,
                  caption="图 4-4 · Round 3 修正：每条 session 算自己的 stats")

    add_para(doc, "**Round 3 结果（8100B / 14 类，最终干净版）：**")
    add_table(
        doc,
        headers=["Method", "论文", "**我们 (Round 3)**", "差距"],
        rows=[
            ["DT", "0.9482", "0.9220", "-0.026"],
            ["RF", "0.9599", "0.9408", "-0.019"],
            ["XGB", "0.9605", "0.9436", "-0.017"],
            ["CNN1D", "0.9609", "0.9406", "-0.020"],
            ["ResNet-18", "0.9773", "0.9454", "-0.032"],
            ["ANDE-no-SE", "0.9726", "**0.9486**", "-0.024"],
            ["**ANDE**", "**0.9820**", "**0.9458**", "**-0.036**"],
        ],
        col_widths_cm=[3.5, 3, 4, 3],
        align=[WD_ALIGN_PARAGRAPH.CENTER] * 4,
    )
    add_para(doc, "所有方法都落在论文报告的 0.94-0.98 区间附近，ANDE 与论文相差 "
                  "~3.6 个百分点。", size=10.5)


def chap_engineering(doc: Document) -> None:
    add_heading(doc, "五、工程上的一个大坑：内存爆炸 → 250 倍提速", level=1)
    add_para(
        doc,
        "Round 3 改成 per-session 算 stats 后，我们立刻撞上工程问题：**单 worker 占 12 GB 内存，"
        "3 小时还跑不完一个 pcap**。",
    )
    add_para(doc, "**初版（错误）实现：**")
    add_code_block(
        doc,
        "sessions = {}\n"
        "for pkt in PcapReader(pcap_path):\n"
        "    key = session_key(pkt)\n"
        "    sessions.setdefault(key, []).append(pkt)   # ← 把整个 scapy 包对象存进 list\n"
        "\n"
        "for key, packets in sessions.items():\n"
        "    feats = compute_features(packets)          # 最后一次性算\n"
        "    rows.append(feats)",
    )
    add_para(doc, "**问题：**")
    add_bullet(doc, "scapy 的 Packet 对象很「重」，每个几 KB")
    add_bullet(doc, "FILE-TRANSFER 类的 pcap 有 ~1600 万 packet，单 worker RSS 飙到 12 GB")
    add_bullet(doc, "bytes(pkt[TCP].payload) 这种调用还会复制 payload，慢上加慢")
    add_bullet(doc, "AutoDL 上跑了 3 小时还没出第一个 pcap")

    add_para(doc, "**修复方案：在线统计（Welford 算法）+ Header 算术**")
    add_code_block(
        doc,
        "class SessionAcc:\n"
        "    __slots__ = ('n', 'pl_mean', 'pl_M2', 'pl_min', 'pl_max', ...)\n"
        "    def update(self, pkt):\n"
        "        self.n += 1\n"
        "        # Welford 在线均值/方差：O(1) 空间\n"
        "        delta = pkt_len - self.pl_mean\n"
        "        self.pl_mean += delta / self.n\n"
        "        self.pl_M2   += delta * (pkt_len - self.pl_mean)\n"
        "        # payload 大小用 IP/TCP 头部算术得到，不复制 payload 字节\n"
        "        ip_payload = ip.len - ip.ihl*4\n"
        "        tcp_payload= ip_payload - tcp.dataofs*4\n"
        "        ...\n"
        "\n"
        "# 主循环\n"
        "session_accs = {}\n"
        "for pkt in PcapReader(pcap_path):\n"
        "    key = session_key(pkt)\n"
        "    session_accs.setdefault(key, SessionAcc()).update(pkt)\n"
        "    # ↑ 这里 update 完，packet 对象就被丢弃了",
    )
    add_callout(
        doc,
        "效果",
        "内存：12 GB / worker  →  几百 MB / worker\n"
        "速度：3 小时跑不完  →  **64 worker × 31 GB pcap = 12 分钟搞定**\n"
        "提速约 250 倍。",
        color_hex="E0F4E0",
    )


def chap_matrix_results(doc: Document) -> None:
    add_heading(doc, "六、完整 42 组复现矩阵", level=1)
    add_para(
        doc,
        "Round 3 干净后，跑完整复现矩阵：3 种长度 × 2 种任务 × 7 种方法 = "
        "**42 个训练运行**。单 seed=42，在 RTX 5090 上 **37 分钟跑完**。",
    )
    if (FIG / "matrix_overview.png").exists():
        add_image(doc, FIG / "matrix_overview.png", width_cm=15,
                  caption="图 6-1 · 42 组实验全景图。红色虚线 = 论文 ANDE 0.9820 目标")

    add_para(doc, "**Behavior14 / 8100B 详细结果：**")
    add_table(
        doc,
        headers=["Method", "Accuracy", "F1", "FPR"],
        rows=[
            ["**ANDE-no-SE**", "**0.9486**", "0.9468", "0.0044"],
            ["ANDE", "0.9458", "0.9454", "0.0046"],
            ["ResNet-18", "0.9454", "0.9439", "0.0047"],
            ["XGB", "0.9436", "0.9423", "0.0051"],
            ["RF", "0.9408", "0.9390", "0.0055"],
            ["CNN1D", "0.9406", "0.9367", "0.0053"],
            ["DT", "0.9220", "0.9222", "0.0068"],
        ],
        col_widths_cm=[4, 3.5, 3.5, 3.5],
        align=[None, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
               WD_ALIGN_PARAGRAPH.CENTER],
    )

    if (FIG / "training_curves.png").exists():
        add_image(doc, FIG / "training_curves.png", width_cm=15,
                  caption="图 6-2 · 8100B / 14 类 ANDE 训练曲线，~20 epoch 后过拟合，early stopping 接住")
    if (FIG / "confusion_matrix.png").exists():
        add_image(doc, FIG / "confusion_matrix.png", width_cm=14,
                  caption="图 6-3 · 8100B / 14 类 ANDE 混淆矩阵。错分集中在同活动的 Tor↔NonTor 之间")

    add_heading(doc, "6.1 SE Block 的消融：和论文方向一致", level=2)
    add_table(
        doc,
        headers=["长度", "ANDE", "ANDE-no-SE", "ΔSE", "论文方向"],
        rows=[
            ["784", "0.9420", "0.9352", "**+0.0068**", "✅ 一致"],
            ["4096", "0.9464", "0.9426", "**+0.0038**", "✅ 一致"],
            ["8100", "0.9458", "0.9486", "-0.0028", "单 seed 噪声"],
        ],
        col_widths_cm=[2.5, 3, 3, 3.5, 4],
        align=[WD_ALIGN_PARAGRAPH.CENTER] * 5,
    )
    add_para(doc, "在 784 / 4096 长度上，SE 提升 0.4-0.7 pp，方向和论文一致。"
                  "8100 略低于无 SE，是单 seed 噪声范畴。", size=10.5)

    add_heading(doc, "6.2 论文复现的方法论价值", level=2)
    add_para(
        doc,
        "我们这次复现的核心贡献，并不是「数字接近论文」，而是 **通过实验暴露并修复了"
        "论文 pseudocode 的歧义**：",
    )
    add_bullet(doc, "字面照搬伪代码 → 跑出 DT/RF/XGB 100% → 触发警报")
    add_bullet(doc, "切到 pcap-level split → 数字断崖式下跌 → 验证泄漏存在")
    add_bullet(doc, "重读 Table II + Algorithm 1 line 5 → 找到论文真实意图")
    add_para(
        doc,
        "这对后续复现者很有用：**不要默认论文的 pseudocode 是字面准确的**，"
        "要用 Table、Section 文字、样本数量去交叉验证。",
    )


def chap_repro_summary(doc: Document) -> None:
    add_heading(doc, "七、复现部分的总结", level=1)

    add_heading(doc, "7.1 三条核心教训", level=2)
    add_bullet(
        doc,
        "**1. 论文 pseudocode 可能有歧义**：Algorithm 2 写 `for *.pcap` 看起来是 per-pcap，"
        "但 Table II 文字明确说 in one session。两者矛盾时要相信样本量、文字描述，"
        "而不是字面伪代码。",
    )
    add_bullet(
        doc,
        "**2. 数据泄漏比想象的更容易发生**：Round 1 的 0.9908 看起来「漂亮」，"
        "真正的报警信号是 DT/RF/XGB 全部 100% 这种不合理的基线表现。"
        "跑全矩阵基线就是为了暴露这种异常。",
    )
    add_bullet(
        doc,
        "**3. 修 bug 时容易过度修正**：Round 2 把 split 改成 pcap-level 确实"
        "堵住了 Round 1 的泄漏，但改变了任务定义，跌到 0.66 完全没法和论文对比。"
        "「症状治好了，病变成了另一种病。」",
    )

    add_heading(doc, "7.2 复现部分的产物", level=2)
    add_bullet(doc, "**代码**：从预处理（preprocess_raw.py + preprocess_stats.py）"
                    "到训练（train.py + evaluate.py）到 7 种方法的完整实现")
    add_bullet(doc, "**实验**：42 组主矩阵 + 三轮历史结果")
    add_bullet(doc, "**最终复现结果**：ANDE 8100B/14 类 = **0.9458**（论文 0.9820）")
    add_bullet(doc, "**仓库**：[github.com/Dwinovo/ANDE-Reproduction](https://github.com/Dwinovo/ANDE-Reproduction)")

    add_callout(
        doc,
        "下一阶段汇报预告",
        "干净复现完之后，我并没有停下，而是继续问了三个论文没回答的问题：\n"
        "  Q1：把字节硬 reshape 成二维方阵真的合理吗？\n"
        "  Q2：固定 784/4096/8100 截断太朴素，能不能自适应选段？\n"
        "  Q3：攻击者用 padding / 延迟 / 流量整形能不能躲过 ANDE？\n"
        "这些会放在「扩展实验」那场汇报里讲。",
        color_hex="EBF5FB",
    )

    add_para(doc, "")
    add_para(doc, "──── 复现部分 END ────", align=WD_ALIGN_PARAGRAPH.CENTER,
             italic=True, size=10)


# =====================================================================
# Extension chapters
# =====================================================================

def recap_for_extension(doc: Document) -> None:
    """Brief 1-page recap of what ANDE is and what we reproduced, so the
    extension doc can stand alone."""
    add_heading(doc, "回顾：第一场汇报讲了什么", level=1)
    add_para(doc, "本场汇报承接上一场「复现部分」。如果你没听过上一场，"
                  "下面这一页就足够你听懂今天的扩展实验：", size=10.5,
             italic=True)

    add_heading(doc, "0.1 ANDE 是干嘛的", level=2)
    add_para(doc,
             "ANDE 是 IEEE TNSM 2024 论文的方法，用来检测匿名网络流量（Tor 等）。"
             "它把一条 session 同时表示成两种形式："
             "**(a)** 把字节流 reshape 成 28×28 / 64×64 / 90×90 的灰度图，"
             "**(b)** 算 26 个手工统计特征（包大小、时延、TCP flag 占比等）。"
             "然后用 SE-ResNet-18（图像分支）+ MLP（特征分支）+ 融合 MLP 做 14 类分类。")

    add_heading(doc, "0.2 14 类任务长什么样", level=2)
    add_para(doc, "7 种行为 × 2 种模式 = 14 类：")
    add_code_block(
        doc,
        "browsing-nontor    chat-nontor    email-nontor    ft-nontor    p2p-nontor    streaming-nontor    voip-nontor\n"
        "browsing-tor       chat-tor       email-tor       ft-tor       p2p-tor       streaming-tor       voip-tor",
    )

    add_heading(doc, "0.3 我们复现到了什么程度", level=2)
    add_table(
        doc,
        headers=["指标", "论文", "**我们干净复现**", "差距"],
        rows=[
            ["Accuracy", "0.9820", "**0.9458**", "-0.036"],
            ["F1", "0.9821", "0.9454", "-0.037"],
            ["FPR", "0.0017", "0.0046", "+0.0029"],
            ["样本数", "50,905", "24,995 sessions", "—"],
            ["训练数据", "33 GB pcap", "33 GB pcap (154 个文件)", "—"],
        ],
        col_widths_cm=[4, 3, 4.5, 3],
        align=[None, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
               WD_ALIGN_PARAGRAPH.CENTER],
    )
    add_para(doc, "和论文差 3.6 pp（百分点），但落在同一量级。"
                  "复现的核心难点是修了一个数据泄漏 bug：论文伪代码字面是"
                  "「per-pcap 算 stats」，但 Table II 文字明确写的是「in one session」。"
                  "前者会让 train/test 共享 stats 导致 DT/RF/XGB 跑到 100% 假象，"
                  "后者才是论文真实意图，这才是我们最终的 0.9458 来源。", size=10.5)

    add_callout(
        doc,
        "本场汇报的起点",
        "我们今天讨论的，全部都是在 **干净的 0.9458 基线** 上做的扩展。"
        "也就是说，所有训练 / 测试都用了「per-session stats + session 级 split」"
        "这个正确口径，没有任何泄漏。",
    )


def chap_ext_intro(doc: Document) -> None:
    add_heading(doc, "一、扩展实验：三个我自己追问的问题", level=1)
    add_para(doc, "论文做完了什么、没做什么？以下是论文没有回答、但我觉得"
                  "**必须问** 的三个问题：")

    add_table(
        doc,
        headers=["#", "我提的问题", "实验设计", "用到的新模型 / 新工具"],
        rows=[
            ["Q1", "把字节排方阵真的合理吗？",
             "做一个 1D CNN，直接处理字节序列，和 ANDE 比",
             "**ByteTCN** (dilated 1-D CNN)"],
            ["Q2", "固定 784/4096/8100 截断太朴素了，能不能自适应选段？",
             "用 Transformer + 注意力让模型挑最关键的 byte 段",
             "**ByteSegmentAttention**"],
            ["Q3", "攻击者用 padding / 延迟 / 流量整形能不能躲过检测？",
             "在测试时给输入加扰动，看 accuracy 掉多少",
             "**input-level attack proxies**"],
        ],
        col_widths_cm=[1, 4.5, 6, 4],
        align=[WD_ALIGN_PARAGRAPH.CENTER, None, None, None],
    )

    add_heading(doc, "1.1 实验组织方式：Phase A → Phase B", level=2)
    add_para(doc, "三个问题都用了「先单 seed 探索，再多 seed 上规模」的方式回答：")
    add_table(
        doc,
        headers=["阶段", "配置", "目的", "实验数"],
        rows=[
            ["**Phase A**", "size=8100, seed=42",
             "快速探索 5 个新模型 + 5 种攻击代理", "7 模型 × 5 条件 = 35"],
            ["**Phase B**", "size=784/4096/8100, seed=42/43/44",
             "上规模验证 ByteTCN 的优势不是单 seed 偶然",
             "3 模型 × 3 size × 3 seed × 5 条件 = 135"],
        ],
        col_widths_cm=[2.5, 4, 5.5, 3.5],
        align=[WD_ALIGN_PARAGRAPH.CENTER, None, None, WD_ALIGN_PARAGRAPH.CENTER],
    )
    add_para(doc, "总共 **~170 次训练运行**，在 RTX 5090 上跑了大约 5 个小时。", size=10.5)


def chap_ext_q1(doc: Document) -> None:
    add_heading(doc, "二、问题 1：把字节排成方阵真的合理吗？", level=1)

    add_heading(doc, "2.1 为什么怀疑？", level=2)
    add_para(doc, "8100 个字节本来是一维有序的，硬 reshape 成 90×90 之后，"
                  "二维卷积会看到这种「伪邻接」：")
    add_code_block(
        doc,
        "原始序列  : byte[0], byte[1], ..., byte[89], byte[90], byte[91], ...\n"
        "90×90 图像:\n"
        "  行 0: byte[ 0]  byte[ 1]  ...  byte[88]  byte[89]\n"
        "  行 1: byte[90]  byte[91]  ...  byte[178] byte[179]\n"
        "                  ↑\n"
        "          二维卷积会把 byte[1] 和 byte[91] 当邻居\n"
        "          但它们在网络协议里其实毫无关系",
    )

    add_heading(doc, "2.2 对照实验：ByteTCN（一维空洞残差卷积网络）", level=2)
    add_para(doc, "ByteTCN 直接把 flatten 后的字节序列当一维信号，"
                  "用 dilation 指数扩张感受野。这样既不引入二维伪邻接，"
                  "又能看到几千 byte 的上下文：")
    add_code_block(
        doc,
        "ByteTCN 结构（src/ande/models/byte_sequence.py）:\n"
        "  Conv1d(1, 96, k=15, stride=2)  → BN → ReLU → MaxPool1d\n"
        "  ┌ ResidualBlock(channels=96, dilation=1)\n"
        "  ├ ResidualBlock(channels=96, dilation=2)\n"
        "  ├ ResidualBlock(channels=96, dilation=4)\n"
        "  └ ResidualBlock(channels=96, dilation=8)\n"
        "  → GlobalPool (mean + max concat) → Linear(192 → 128) → Linear(128 → 14)",
    )

    add_heading(doc, "2.3 Phase A 单 seed 结果", level=2)
    if (FIG / "experiment_process_phaseA_clean.png").exists():
        add_image(doc, FIG / "experiment_process_phaseA_clean.png", width_cm=15,
                  caption="图 2-1 · Phase A 干净测试集上 7 个模型的对比")

    add_table(
        doc,
        headers=["Method", "Accuracy", "F1", "FPR", "训练时间 (s)"],
        rows=[
            ["**ByteTCN**", "**0.9568**", "**0.9531**", "**0.0035**", "198.3"],
            ["ByteTCN + stats", "0.9526", "0.9478", "0.0039", "152.1"],
            ["ANDE", "0.9458", "0.9454", "0.0046", "70.7"],
            ["ResNet-18", "0.9454", "0.9439", "0.0047", "77.5"],
            ["CNN1D（baseline）", "0.9406", "0.9367", "0.0053", "100.5"],
            ["SegmentAttention + stats", "0.9370", "0.9360", "0.0052", "65.4"],
            ["SegmentAttention", "0.9364", "0.9329", "0.0056", "42.7"],
        ],
        col_widths_cm=[5, 2.5, 2.5, 2.5, 2.5],
        align=[None] + [WD_ALIGN_PARAGRAPH.CENTER] * 4,
    )
    add_para(doc, "**ByteTCN 在 clean test 上比 ANDE 高出 1.1 个百分点。** "
                  "更重要的是，这不是「随便一个一维模型都行」——简单的 CNN1D（论文里的 baseline）"
                  "只有 0.9406，反而比 ANDE 弱。说明赢的是「**强大的一维序列模型**」，"
                  "不是「**一维**」本身。")

    add_heading(doc, "2.4 Phase B 三 seed × 三长度 上规模验证", level=2)
    add_para(doc, "为了排除单 seed 偶然，跑了 3 个 seed (42/43/44) × 3 种长度的扩大实验：")

    if (FIG / "experiment_process_phaseB_clean.png").exists():
        add_image(doc, FIG / "experiment_process_phaseB_clean.png", width_cm=15,
                  caption="图 2-2 · Phase B：3 种长度 × 3 个 seed 下 ByteTCN vs ANDE")

    add_table(
        doc,
        headers=["长度", "ANDE 平均", "ByteTCN 平均", "ByteTCN 增益"],
        rows=[
            ["784", "0.9407 ± 0.0019", "**0.9547 ± 0.0024**",
             "**+1.40 pp ± 0.30**"],
            ["4096", "0.9447 ± 0.0013", "0.9533 ± 0.0014",
             "+0.86 pp ± 0.07"],
            ["8100", "0.9443 ± 0.0019", "**0.9560 ± 0.0012**",
             "**+1.17 pp ± 0.07**"],
        ],
        col_widths_cm=[2.5, 4, 4, 4],
        align=[WD_ALIGN_PARAGRAPH.CENTER] * 4,
    )

    add_callout(
        doc,
        "Q1 结论",
        "ByteTCN 在 3 个长度、3 个 seed 上都稳定击败 ANDE，平均高 1 个百分点。\n"
        "更有意思的是：**784B 的 ByteTCN（0.9547）已经超过 8100B 的 ANDE（0.9443）。**\n"
        "→ 模型结构比单纯加长输入更重要。\n"
        "→ 把字节当二维图来卷积，确实是个不合适的归纳偏置。",
        color_hex="E0F4E0",
    )


def chap_ext_q2(doc: Document) -> None:
    add_heading(doc, "三、问题 2：自适应选段能不能比固定截断更好？", level=1)

    add_heading(doc, "3.1 为什么怀疑？", level=2)
    add_para(doc, "论文用 784 / 4096 / 8100 这种固定长度截断，"
                  "等于默认「session 前 N 个字节最重要」。但不同 session 的关键信息"
                  "可能出现在不同位置——HTTPS handshake 在最前面，"
                  "TLS application data 在中间，关闭信号在后面。")

    add_heading(doc, "3.2 我们的设计：ByteSegmentAttention", level=2)
    add_code_block(
        doc,
        "ByteSegmentAttention 流程：\n"
        "  1) 把 8100 字节切成 64 个不重叠的 128-byte segment\n"
        "  2) Linear(128 → 128) + 位置编码，得到 (B, 64, 128)\n"
        "  3) 2 层 Transformer encoder 给 segment 之间建立上下文\n"
        "  4) Attention pooling：用一个评分网络给每个 segment 打分，softmax 后加权求和\n"
        "  5) Linear → 14 类",
    )
    add_para(doc, "**重点是步骤 4 的 attention score**——它告诉我们「模型认为哪些"
                  "byte 段最有信息量」。", size=10.5)

    add_heading(doc, "3.3 模型学到了什么？", level=2)
    if (FIG / "experiment_process_attention_segments.png").exists():
        add_image(doc, FIG / "experiment_process_attention_segments.png", width_cm=15,
                  caption="图 3-1 · SegmentAttention 学到的 Top byte ranges（标准化后均值）")

    add_para(doc, "**学到的 top 5 byte ranges：**")
    add_table(
        doc,
        headers=["模型", "Top 5 byte 范围"],
        rows=[
            ["SegmentAttention（纯 byte）",
             "0-128, 128-256, 768-896, 384-512, 1024-1152"],
            ["SegmentAttention + stats",
             "128-256, 8064-8192, 0-128, 384-512, 768-896"],
        ],
        col_widths_cm=[6, 9.5],
        align=[None, None],
    )

    add_para(doc, "**注意 attention 集中在 session 开头**——和直觉一致："
                  "TLS handshake、HTTP header 都在前几百字节。"
                  "「+ stats」版本还会回头看一下 8064-8192（接近 8100 的尾部），"
                  "可能是 stats 提供了「这条 session 多长」的提示。")

    add_heading(doc, "3.4 但是—— accuracy 没超过 ANDE", level=2)
    add_table(
        doc,
        headers=["Method", "Accuracy", "F1", "评价"],
        rows=[
            ["ANDE", "0.9458", "0.9454", "复现基线"],
            ["SegmentAttention + stats", "0.9370", "0.9360", "落后 0.9 pp"],
            ["SegmentAttention（纯 byte）", "0.9364", "0.9329", "落后 0.9 pp"],
        ],
        col_widths_cm=[5.5, 3, 3, 4],
        align=[None, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, None],
    )

    add_callout(
        doc,
        "Q2 结论",
        "自适应选段的方向是对的——attention 没有均匀撒在整个 session 上，"
        "而是集中在 session 开头几百字节。这个分布符合协议直觉。\n"
        "但 accuracy 没超 ANDE，说明当前模型太浅、segment 切得太粗。\n"
        "下一步候选：① TCN 编码器 + top-k 池化  ② 重叠 segment  ③ 多尺度 segment。",
        color_hex="FFF4DC",
    )


def chap_ext_q3(doc: Document) -> None:
    add_heading(doc, "四、问题 3：攻击者改造流量能躲开检测吗？", level=1)
    add_para(doc, "**这才是这个研究的现实意义**——如果攻击者知道你在用 ANDE 检测，"
                  "他完全可以反过来给自己的 Tor 流量加干扰。常见的反制手段有三类：")

    add_table(
        doc,
        headers=["攻击手段", "现实操作", "我们的代理 (proxy)"],
        rows=[
            ["填充（padding）",
             "在真实包前后塞一些无意义的「掩护」字节",
             "**random_padding**：在 image 前 15% 插入随机 byte，截掉尾部"],
            ["随机延迟（delay）",
             "在每个包发送前等一个随机时间",
             "**random_delay**：修改 stats 中 Duration / 各种 delta 时间特征"],
            ["流量整形（shaping）",
             "把所有包压成统一长度（如全部 1500B 或全部 64B）",
             "**traffic_shaping**：缩小 stats 中 packet length / payload 系列特征"],
            ["组合",
             "上述三种一起用",
             "**combined** = padding + delay + shaping"],
        ],
        col_widths_cm=[3, 5.5, 7],
        align=[None, None, None],
    )

    add_para(
        doc,
        "**注意：** 这些是测试时（test-time）对模型输入做的扰动，不是真改 pcap。"
        "属于「rapid prototyping」——能快速看模型对哪种攻击形态最敏感，但不能直接"
        "等价于真实攻击成功率。",
        size=10.5,
    )

    add_para(doc, "**代码片段（src/ande/attacks.py）：**")
    add_code_block(
        doc,
        "def perturb_image(image, spec):\n"
        "    if spec.name == 'random_padding':\n"
        "        # 把前 15% 替换成随机 byte，原内容向后移、超出部分丢弃\n"
        "        flat = image.flatten()\n"
        "        n = int(0.15 * flat.numel())\n"
        "        prefix = torch.rand(n, dtype=flat.dtype)\n"
        "        return torch.cat([prefix, flat[:-n]]).view_as(image)\n"
        "\n"
        "def perturb_stat(stat, spec):\n"
        "    if spec.name == 'random_delay':\n"
        "        # 把所有时间相关的特征 +0.75（标准化空间里）\n"
        "        for feat in ['Duration_window_flow', 'Avg_deltas_time',\n"
        "                     'Min_deltas_time', 'Max_deltas_time']:\n"
        "            stat[FEATURE_INDEX[feat]] += 0.75",
    )

    add_heading(doc, "4.1 Phase A 鲁棒性结果", level=2)
    if (FIG / "experiment_process_phaseA_padding.png").exists():
        add_image(doc, FIG / "experiment_process_phaseA_padding.png", width_cm=15,
                  caption="图 4-1 · Phase A：5 种条件下 7 个模型的 accuracy")

    add_table(
        doc,
        headers=["Method", "clean", "padding", "delay", "shaping", "combined"],
        rows=[
            ["**ByteTCN**", "**0.9568**", "**0.8018**", "0.9568",
             "0.9568", "**0.8020**"],
            ["ByteTCN + stats", "0.9526", "0.8108", "0.9526",
             "0.9524", "0.8110"],
            ["**ANDE**", "**0.9458**", "**0.1026**", "0.9456",
             "0.9458", "**0.0990**"],
            ["ResNet-18", "0.9454", "0.2511", "0.9454", "0.9454", "0.2531"],
            ["CNN1D", "0.9406", "0.2679", "0.9406", "0.9406", "0.2659"],
            ["SegmentAttention + stats", "0.9370", "0.2068", "0.9364",
             "0.9360", "0.2016"],
            ["SegmentAttention", "0.9364", "0.1140", "0.9364",
             "0.9364", "0.1160"],
        ],
        col_widths_cm=[5, 2.2, 2.2, 2.2, 2.2, 2.2],
        align=[None] + [WD_ALIGN_PARAGRAPH.CENTER] * 5,
        size=9,
    )

    add_para(doc, "**三个非常关键的现象：**")
    add_bullet(
        doc,
        "**delay 和 shaping 几乎没杀伤力**：所有模型 accuracy 基本不变。"
        "说明模型其实主要靠 raw bytes 决策，stats 那一路对最终输出影响很小。",
    )
    add_bullet(
        doc,
        "**padding 是最致命的攻击**：ANDE 从 0.9458 直接掉到 **0.1026**（!），"
        "几乎和瞎猜一样（14 类瞎猜 = 1/14 ≈ 0.071）。",
    )
    add_bullet(
        doc,
        "**ByteTCN 对 padding 远更鲁棒**：只从 0.9568 掉到 0.8018，"
        "仍然能正确识别 80% 的 session。",
    )

    add_heading(doc, "4.2 Phase B 上规模验证", level=2)
    if (FIG / "experiment_process_phaseB_padding_drop.png").exists():
        add_image(doc, FIG / "experiment_process_phaseB_padding_drop.png", width_cm=15,
                  caption="图 4-2 · Phase B：random padding 让 accuracy 下降多少")

    add_table(
        doc,
        headers=["长度", "Method", "clean", "padded", "drop"],
        rows=[
            ["784", "ANDE", "0.9407", "0.1174", "**-0.8234**"],
            ["784", "ByteTCN", "0.9547", "0.8598", "**-0.0949**"],
            ["4096", "ANDE", "0.9447", "0.1617", "**-0.7830**"],
            ["4096", "ByteTCN", "0.9533", "0.9416", "**-0.0117**"],
            ["8100", "ANDE", "0.9443", "0.0743", "**-0.8699**"],
            ["8100", "ByteTCN", "0.9560", "0.8022", "**-0.1538**"],
        ],
        col_widths_cm=[2, 4, 2.7, 2.7, 3.5],
        align=[WD_ALIGN_PARAGRAPH.CENTER, None,
               WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
               WD_ALIGN_PARAGRAPH.CENTER],
    )

    add_callout(
        doc,
        "Q3 结论",
        "在 3 种长度 × 3 个 seed 的稳定实验下：\n"
        "  • ANDE 在 prefix padding 下平均掉 78-87 pp，几乎完全失效\n"
        "  • ByteTCN 平均只掉 1-15 pp，4096B 时甚至几乎不受影响\n"
        "  • delay / shaping proxy 影响很小，说明 stats 分支决策权重很低\n"
        "→ **现实启示：如果要部署这类匿名流量检测器，必须考虑「攻击者会前缀填充」的威胁模型，"
        "并优先采用 ByteTCN 这种序列模型。**",
        color_hex="E0F4E0",
    )


def chap_ext_summary(doc: Document) -> None:
    add_heading(doc, "五、扩展实验的整体总结与下一步", level=1)

    add_heading(doc, "5.1 三条核心洞察", level=2)
    add_bullet(
        doc,
        "**1. 一维序列模型 > 二维方阵模型**：ByteTCN 在 3 个长度、3 个 seed 上"
        "稳定击败 ANDE。说明 reshape 引入的二维邻接是错误归纳偏置。",
    )
    add_bullet(
        doc,
        "**2. 自适应选段方向是对的，但需要更强模型**：SegmentAttention 还没超过 ANDE，"
        "但 attention 集中在 session 开头几百字节——和 TLS handshake / HTTP header "
        "的位置吻合，说明信号确实在那里。",
    )
    add_bullet(
        doc,
        "**3. ANDE 对前缀填充攻击非常脆弱**：87 pp drop 不是噪声，3 个 seed 都稳定复现。"
        "现实部署里这是个严重问题。",
    )

    add_heading(doc, "5.2 扩展实验产物清单", level=2)
    add_bullet(doc, "**新代码**：byte_sequence.py (ByteTCN + ByteSegmentAttention)、"
                    "attacks.py、Phase A/B 两套 orchestrator 脚本")
    add_bullet(doc, "**实验数**：Phase A 35 行 + Phase B 135 行 = **170 个训练运行**")
    add_bullet(doc, "**完整结果文件**：docs/results/extended_phaseA.csv、"
                    "extended_phaseB.csv、extended_phaseB.summary.json")
    add_bullet(doc, "**计算量**：RTX 5090 上 ~5 小时")

    add_heading(doc, "5.3 下一步可以做什么", level=2)
    add_bullet(doc, "**真实 pcap 级攻击**：现在的 padding/delay 是 input-level proxy，"
                    "需要实现 packet-level 重写 pcap 后再过完整管道验证")
    add_bullet(doc, "**对抗训练**：在训练时也加随机 padding，看 ANDE 能不能恢复鲁棒性")
    add_bullet(doc, "**更强的 SegmentAttention**：换 TCN encoder + top-k pooling，"
                    "预期能超过 ByteTCN")
    add_bullet(doc, "**SOTA 横向对比**：实现 FlowPic / MSerNetDroid 这些论文也提到的"
                    "对照方法，把 ByteTCN 的优势放到更大的对比里验证")

    add_callout(
        doc,
        "一句话收尾",
        "复现的目的不仅是「跑出论文的数字」。在干净复现之上继续追问、做对照实验、"
        "做对抗鲁棒性，才能真正理解这个方法的边界。\n"
        "本场汇报的最大收获：「把字节当一维序列处理」+「警惕前缀填充攻击」"
        "是这类系统在实际部署前必须正视的两件事。",
        color_hex="EBF5FB",
    )

    add_para(doc, "")
    add_para(doc, "──── 扩展实验部分 END ────",
             align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=10)


# =====================================================================
# Build orchestration
# =====================================================================

def build_reproduction() -> None:
    doc = Document()
    _make_normal_style(doc)

    cover_repro(doc)
    chap_repro_intro(doc)
    chap_sample_definition(doc)
    chap_paper_method(doc)
    chap_three_rounds(doc)
    chap_engineering(doc)
    chap_matrix_results(doc)
    chap_repro_summary(doc)

    doc.save(str(OUTPUT_REPRO))
    print(f"Saved: {OUTPUT_REPRO}")


def build_extension() -> None:
    doc = Document()
    _make_normal_style(doc)

    cover_extension(doc)
    recap_for_extension(doc)
    chap_ext_intro(doc)
    chap_ext_q1(doc)
    chap_ext_q2(doc)
    chap_ext_q3(doc)
    chap_ext_summary(doc)

    doc.save(str(OUTPUT_EXT))
    print(f"Saved: {OUTPUT_EXT}")


if __name__ == "__main__":
    build_reproduction()
    build_extension()
